# Builds setup-checklist.docx from the same content as setup-checklist.html.
# Target: one US Letter page, printable, with tick-boxes students fill in by hand.
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK, ACCENT, MUTED = "1F1D1A", "B14A2E", "5B564C"
CREAM, ORANGE, GREEN, RULE = "F0EEE5", "FDE0D2", "D5E6DC", "D8D4C8"
SERIF, SANS, MONO = "Georgia", "Arial", "Consolas"

# --- OOXML child order is schema-enforced: insert into the right slot, never append.
PPR_ORDER = ["w:pStyle","w:keepNext","w:keepLines","w:pageBreakBefore","w:framePr",
    "w:widowControl","w:numPr","w:suppressLineNumbers","w:pBdr","w:shd","w:tabs",
    "w:suppressAutoHyphens","w:kinsoku","w:wordWrap","w:overflowPunct","w:topLinePunct",
    "w:autoSpaceDE","w:autoSpaceDN","w:bidi","w:adjustRightInd","w:snapToGrid","w:spacing",
    "w:ind","w:contextualSpacing","w:mirrorIndents","w:suppressOverlap","w:jc",
    "w:textDirection","w:textAlignment","w:textboxTightWrap","w:outlineLvl","w:divId",
    "w:cnfStyle","w:rPr","w:sectPr","w:pPrChange"]
TCPR_ORDER = ["w:cnfStyle","w:tcW","w:gridSpan","w:hMerge","w:vMerge","w:tcBorders","w:shd",
    "w:noWrap","w:tcMar","w:textDirection","w:tcFitText","w:vAlign","w:hideMark",
    "w:cellIns","w:cellDel","w:cellMerge"]
EDGE_ORDER = ["top","left","bottom","right","between","bar"]

def el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), v)
    return e

def _place(pr, child, order):
    """Insert child into pr at its schema position, replacing any existing one."""
    tag = child.tag.split("}")[1]
    existing = pr.find(qn("w:" + tag))
    if existing is not None:
        pr.remove(existing)
    idx = order.index("w:" + tag)
    for e in pr:
        name = "w:" + e.tag.split("}")[1]
        if name in order and order.index(name) > idx:
            e.addprevious(child)
            return child
    pr.append(child)
    return child

def shade(pr, fill, order):
    _place(pr, el("w:shd", val="clear", color="auto", fill=fill), order)

def set_borders(pr, edges, kind):
    """edges: {edge: (color, sz, space)}. Children of the border element are ordered too."""
    tag, order = ("w:pBdr", PPR_ORDER) if kind == "p" else ("w:tcBorders", TCPR_ORDER)
    b = OxmlElement(tag)
    for e in EDGE_ORDER:
        if e in edges:
            color, sz, space = edges[e]
            b.append(el("w:" + e, val="single", sz=sz, space=space, color=color))
    _place(pr, b, order)

def spacing(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line: pf.line_spacing = line

def run(p, text, size=8.5, font=SERIF, bold=False, italic=False, color=INK):
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    # east-asian font must be set too or Word substitutes
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return r

def para(container, size=8.5, before=0, after=1.5, **kw):
    p = container.add_paragraph()
    spacing(p, before, after, kw.get("line"))
    return p

def code(container, lines, size=7.6):
    """One shaded paragraph per code block, a line break between lines."""
    p = container.add_paragraph()
    spacing(p, 2, 3, 1.0)
    pr = p._p.get_or_add_pPr()
    shade(pr, CREAM, PPR_ORDER)
    set_borders(pr, {"left": (ACCENT, "18", "4"), "top": (CREAM, "2", "2"),
                     "bottom": (CREAM, "2", "2"), "right": (CREAM, "2", "2")}, "p")
    pf = p.paragraph_format
    pf.left_indent = Inches(0.07); pf.right_indent = Inches(0.05)
    for i, ln in enumerate(lines):
        if i: p.add_run().add_break()
        run(p, ln, size=size, font=MONO)
    return p

def onecell(doc, width_in, fill, bcolor, bsz="12"):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    c = t.rows[0].cells[0]
    c.width = Inches(width_in)
    tcPr = c._tc.get_or_add_tcPr()
    set_borders(tcPr, {e: (bcolor, bsz, "0") for e in ("top", "left", "bottom", "right")}, "tc")
    shade(tcPr, fill, TCPR_ORDER)
    return t, c

def cellpad(cell, tw=90):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ("top", "start", "bottom", "end"):
        mar.append(el("w:" + side, w=str(tw), type="dxa"))
    _place(tcPr, mar, TCPR_ORDER)

def code_rich(container, lines, size=7.8):
    """Like code(), but each line is a list of (text, color) pairs."""
    p = container.add_paragraph()
    spacing(p, 2, 3, 1.0)
    pr = p._p.get_or_add_pPr()
    shade(pr, CREAM, PPR_ORDER)
    set_borders(pr, {"left": (ACCENT, "18", "4"), "top": (CREAM, "2", "2"),
                     "bottom": (CREAM, "2", "2"), "right": (CREAM, "2", "2")}, "p")
    p.paragraph_format.left_indent = Inches(0.07)
    for i, parts in enumerate(lines):
        if i: p.add_run().add_break()
        for text, col in parts:
            run(p, text, size=size, font=MONO, color=col)
    return p

# ---------------------------------------------------------------- document
doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(s, m, Inches(0.4))
CW = 7.7

st = doc.styles["Normal"]
st.font.name = SERIF; st.font.size = Pt(9.2)
st._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
st.paragraph_format.space_before = Pt(0); st.paragraph_format.space_after = Pt(1.5)
st.paragraph_format.line_spacing = 1.0

M  = {"font": MONO, "size": 8.2}
Ms = {"font": MONO, "size": 7.9}
Mb = {"font": MONO, "size": 8.2, "bold": True}
B  = {"bold": True}
I  = {"italic": True}

def step(n, title):
    p = para(doc, before=3, after=2)
    pr = p._p.get_or_add_pPr(); set_borders(pr, {"bottom": (RULE, "6", "2")}, "p")
    run(p, f"{n}   ", size=10.5, font=SANS, bold=True, color=ACCENT)
    run(p, title, size=10.5, font=SANS, bold=True)

def tick(runs, size=8.6):
    p = para(doc, after=1.5)
    run(p, "☐  ", size=10, font=SANS)
    for t_, kw in runs:
        kw = dict(kw); kw.setdefault("size", size)
        run(p, t_, **kw)

def rich(container, runs, size=9.2, before=0, after=1.5, left=0.0):
    p = container.add_paragraph(); spacing(p, before, after)
    if left: p.paragraph_format.left_indent = Inches(left)
    for t_, kw in runs:
        kw = dict(kw); kw.setdefault("size", size)
        run(p, t_, **kw)
    return p

def option(label, note=None):
    p = para(doc, before=3, after=1)
    p.paragraph_format.left_indent = Inches(0.08)
    pr = p._p.get_or_add_pPr(); set_borders(pr, {"left": (RULE, "12", "6")}, "p")
    run(p, label, size=8.6, font=SANS, bold=True)
    if note:
        run(p, "   " + note, size=7.5, font=SANS, bold=True, color=ACCENT)

# ---- header
p = para(doc, after=0)
run(p, "DS 5030 · UNDERSTANDING UNCERTAINTY · TUESDAY, AUGUST 25, 2026",
    size=7.5, font=SANS, bold=True, color=ACCENT)
p = para(doc, after=3)
run(p, "Day One Setup Checklist", size=15, font=SANS, bold=True)

t = doc.add_table(rows=1, cols=3); t.autofit = False
for cell, label in zip(t.rows[0].cells, ("Name", "GitHub username", "Operating system")):
    cell.width = Inches(CW / 3)
    set_borders(cell._tc.get_or_add_tcPr(), {"bottom": (INK, "6", "0")}, "tc")
    pp = cell.paragraphs[0]; spacing(pp, 0, 2)
    run(pp, label, size=8, font=SANS, color=MUTED)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---- goal box
t, c = onecell(doc, CW, ORANGE, ACCENT, bsz="14")
cellpad(c, 90)
pp = c.paragraphs[0]; spacing(pp, 0, 1)
run(pp, "Two things finish this: ", size=9.2, bold=True)
run(pp, "SETUP CODE: UU-92-5650", size=9.2, font=MONO, bold=True)
run(pp, " on your screen, and your commit in your own fork.", size=9.2, bold=True)
pp2 = c.add_paragraph(); spacing(pp2, 0, 0)
run(pp2, "Everyone gets the same setup code. Show it, and your fork on github.com, to the instructor "
         "or an IA before you leave. Nothing is graded today, but Thursday's lab assumes all of it works.",
    size=8.8)
doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ---- 1
step(1, "Make a folder for this class")
rich(doc, [("Somewhere you will find again: Documents, not Downloads. Make a course folder with an empty ", {}),
           ("data", M), (" folder inside it. By the end of step 5:", {})])
code_rich(doc, [
  [("DS5030/", INK)],
  [("├── data/       ", INK), ("← datasets you download later in the semester", MUTED)],
  [("└── UU_F26/     ", INK), ("← step 5 puts your fork here", MUTED)],
], size=7.6)
tick([("My class folder exists, and it has an empty ", {}), ("data", Ms), (" folder inside it.", {})])

# ---- 2
step(2, "Open the folder in VS Code, and add two extensions")
rich(doc, [("File > Open Folder", B), (", and choose the folder you just made. Then open Extensions (", {}),
           ("Cmd/Ctrl + Shift + X", M), (") and install both (published by Microsoft):", {})])
tick([("Python", B), (" installed        ", {}), ("☐  ", {"size": 10, "font": SANS}),
      ("Jupyter", B), (" installed", {})])
tick([("The VS Code sidebar shows my class folder, with ", {}), ("data", Ms), (" inside it.", {})])

# ---- 3
step(3, "Create the conda environment")
rich(doc, [("An ", {}), ("environment", I),
           (" is a private Python, so this class cannot break your other projects. No conda yet? Install ", {}),
           ("Miniconda", B), (" from ", {}), ("docs.conda.io/en/latest/miniconda.html", M), (", then ", {}),
           ("close and reopen your terminal", B), (". Run these two lines, one at a time:", {})])
code(doc, [
 'conda create -n uu -c conda-forge python=3.12 numpy "pandas>=2.2,<3" matplotlib \\',
 '      seaborn scipy statsmodels networkx ipykernel jupyterlab',
 'conda activate uu',
], size=7.6)
tick([("It finished without an error (a few minutes), and my prompt now starts with ", {}), ("(uu)", Ms), (".", {})])

# ---- 4
step(4, "Point VS Code at that environment")
rich(doc, [("VS Code will not find it on its own, and this is the step people skip. Press ", {}),
           ("Cmd/Ctrl + Shift + P", M), (", type ", {}), ("Python: Select Interpreter", B),
           (", and pick the one labelled ", {}), ("uu", M), (".", {})])
tick([("The status bar at the bottom of the VS Code window now shows ", {}), ("uu", Ms), (".", {})])

# ---- 5
step(5, "Fork the repository, then clone your fork")
rich(doc, [("You work in your own copy, not mine.", B), (" A ", {}), ("fork", I),
           (" is your own GitHub copy of my repository: you push your work to yours, and pull my updates "
            "from mine. One fresh repository per lab.", {})])
rich(doc, [("Fork it on the web first.", B), (" At ", {}), ("github.com/smgroves/UU_F26", M),
           (", click ", {}), ("Fork", B), (" (top right) and keep the name. You now own ", {}),
           ("github.com/YOUR-USERNAME/UU_F26", M), (".", {})])

option("Option A · GitHub Desktop", "RECOMMENDED")
rich(doc, [("Install from ", {}), ("desktop.github.com", M),
           (" and sign in, which also settles your push permissions. Then ", {}),
           ("File > Clone repository > GitHub.com", B), (", pick the ", {}), ("UU_F26", M),
           (" under ", {}), ("your own", B), (" username, and set the local path to your class folder. ", {}),
           ("If it asks how you plan to use the fork, choose the option about your own purposes", B),
           (", not contributing to the parent: that keeps your pushes going to your fork, not mine.", {})],
     left=0.08)

option("Option B · Command line")
rich(doc, [("In VS Code's terminal (", {}), ("Terminal > New Terminal", B),
           ("), from your class folder, with ", {}), ("your own username", B), (":", {})], left=0.08)
code(doc, ["git clone https://github.com/YOUR-USERNAME/UU_F26.git"], size=7.6)

rich(doc, [("Then everyone, in VS Code's terminal", B), (", add my copy as ", {}), ("upstream", M),
           (" and check both remotes:", {})], before=3)
code(doc, ["cd UU_F26",
           "git remote add upstream https://github.com/smgroves/UU_F26.git",
           "git remote -v"], size=7.6)
tick([("origin", Ms), (" shows ", {}), ("my own username", B), (" (where my work goes) and ", {}),
      ("upstream", Ms), (" shows ", {}), ("smgroves", Ms), (" (where updates come from).", {})])
tick([("UU_F26", Ms), (" is in the VS Code sidebar next to ", {}), ("data", Ms), (", and contains ", {}),
      ("Day_1_python_test.ipynb", Ms), (".", {})])

# ================================================================ page 2
p = doc.add_paragraph(); spacing(p, 0, 0)
p.add_run().add_break(WD_BREAK.PAGE)
p = para(doc, after=7)
pr = p._p.get_or_add_pPr(); set_borders(pr, {"bottom": (RULE, "6", "3")}, "p")
run(p, "DS 5030 · Day One Setup · page 2 of 2 — run the test, push your work, then fix whatever broke",
    size=8, font=SANS, bold=True, color=MUTED)

# ---- 6
step(6, "Run the notebook, and check every line")
rich(doc, [("Open ", {}), ("UU_F26/Day_1_python_test.ipynb", M), (" from the sidebar. ", {}),
           ("Look at the top right: the kernel must say ", B), ("uu", Mb), (".", B), (" If it says ", {}),
           ("base", M), (", ", {}), ("Python 3", M), (", or ", {}), ("Select Kernel", M),
           (", click it and pick ", {}), ("uu", M), (". Then ", {}), ("Run All", B),
           (", and tick a box for each package that printed a version number:", {})])
p = para(doc, after=2)
for i, pkg in enumerate(("numpy", "pandas", "matplotlib", "seaborn", "scipy", "statsmodels")):
    if i: run(p, "      ", size=8.6)
    run(p, "☐ ", size=10, font=SANS)
    run(p, pkg, size=8.4, font=MONO)
tick([("The ", {}), ("interpreter", Ms), (" line contains ", {}), ("/uu/", Ms),
      (". If not, the kernel is wrong: fix it before going on.", {})])
tick([("A histogram appeared.", B), (" Not a warning, not a blank space: an actual picture.", {})])
tick([("The last line reads exactly ", {}), ("SETUP CODE: UU-92-5650", Ms), (".", {})])

# ---- 7
step(7, "Push the finished notebook to your fork")
rich(doc, [("Save the notebook first (", {}), ("Cmd/Ctrl + S", M),
           ("), so the run you just did is on disk.", {})])
option("Option A · GitHub Desktop", "RECOMMENDED")
rich(doc, [("It already lists ", {}), ("Day_1_python_test.ipynb", M), (" as changed. Type a summary such as ", {}),
           ("setup complete", M), (", click ", {}), ("Commit to main", B), (", then ", {}),
           ("Push origin", B), (".", {})], left=0.08)
option("Option B · Command line")
code(doc, ["git add Day_1_python_test.ipynb",
           'git commit -m "setup complete"',
           "git push"], size=7.6)
tick([("I opened ", {}), ("github.com/YOUR-USERNAME/UU_F26", Ms), (" in a browser and ", {}),
      ("my commit is there", B), (".", {})])
tick([("I showed that page, and the setup code, to the instructor or an IA.", {})])
rich(doc, [("Pulling my updates later.", B), (" Easiest route: click ", {}), ("Sync fork", B),
           (" on your fork's github.com page, then pull. From the terminal: ", {}),
           ("git pull upstream main", M), (". Do this before every class. ", {}),
           ("Next week's lab is group work", B), (", and we will add collaboration then.", {})],
     before=4, after=5)

# ---- asides
t = doc.add_table(rows=1, cols=2); t.autofit = False
left, right = t.rows[0].cells
for cell, fill, bcol in ((left, "FFFFFF", RULE), (right, GREEN, "7FA892")):
    cell.width = Inches(CW / 2 - 0.06)
    tcPr = cell._tc.get_or_add_tcPr()
    set_borders(tcPr, {e: (bcol, "6", "0") for e in ("top", "left", "bottom", "right")}, "tc")
    shade(tcPr, fill, TCPR_ORDER)
    cellpad(cell, 95)

def ahead(cell, text):
    pp = cell.paragraphs[0]; spacing(pp, 0, 3)
    run(pp, text, size=9.2, font=SANS, bold=True)

def ap(cell, runs, size=8.5):
    pp = cell.add_paragraph(); spacing(pp, 0, 2.5)
    for t_, kw in runs:
        kw = dict(kw); kw.setdefault("size", size)
        run(pp, t_, **kw)
    return pp

AMs = {"font": MONO, "size": 8.0}
AMb = {"font": MONO, "size": 8.0, "bold": True}

ahead(left, "If it breaks")
ap(left, [('"conda: command not found"', AMb),
          (": you did not reopen the terminal after installing Miniconda. On Windows, use the ", {}),
          ("Anaconda Prompt", B), (" rather than the default terminal.", {})])
ap(left, [("The kernel picker does not offer ", B), ("uu", AMb),
          (": quit VS Code entirely and reopen. Still missing? You left ", {}), ("ipykernel", AMs),
          (" out of step 3: run ", {}), ("conda install -n uu ipykernel", AMs), (".", {})])
ap(left, [("No ", B), ("/uu/", AMb), (" in the interpreter line", B),
          (": the notebook is running a different Python than your packages went into. ", {}),
          ("The most common problem in this class", B),
          (", and fixing it now beats fixing it Thursday. Redo step 4, then re-pick the kernel.", {})])
ap(left, [("origin", AMb), (" says ", B), ("smgroves", AMb),
          (": you cloned mine instead of your fork. Repoint it with ", {}),
          ("git remote set-url origin https://github.com/YOUR-USERNAME/UU_F26.git", AMs), (".", {})])
ap(left, [('"git: command not found"', AMb),
          (": install GitHub Desktop, which includes git. On a new Mac you can instead run ", {}),
          ("xcode-select --install", AMs), (" and wait for it to finish.", {})])
ap(left, [("The push is rejected, or asks for a password", B),
          (": GitHub stopped accepting passwords in 2021. Use GitHub Desktop, which handles this for you, "
           "or set up a personal access token or SSH key. Come find us rather than fighting it.", {})])
ap(left, [('"No such file or directory: cville_cars.csv"', AMb),
          (": the notebook has to run from inside ", {}), ("UU_F26", AMs),
          (". Open it from the sidebar rather than moving the file.", {})])

ahead(right, "Fallback: Google Colab")
ap(right, [("If your machine will not install anything, you can still prove the code works in the browser. "
            "Go to ", {}), ("colab.research.google.com", AMs), (", then ", {}),
           ("File > Open notebook > GitHub", B), (", and paste your own fork:", {})])
for lines in (["YOUR-USERNAME/UU_F26"],
              ["!git clone https://github.com/YOUR-USERNAME/UU_F26.git", "%cd UU_F26"]):
    if lines[0].startswith("!"):
        ap(right, [("Open ", {}), ("Day_1_python_test.ipynb", AMs), (", add ", {}),
                   ("one new cell at the very top", B), (", and run it first so the data file is there:", {})])
    pp = right.add_paragraph(); spacing(pp, 2, 3, 1.0)
    pr = pp._p.get_or_add_pPr(); shade(pr, "FFFFFF", PPR_ORDER)
    set_borders(pr, {"left": ("7FA892", "18", "4")}, "p")
    pp.paragraph_format.left_indent = Inches(0.06)
    for i, ln in enumerate(lines):
        if i: pp.add_run().add_break()
        run(pp, ln, size=7.6, font=MONO)
ap(right, [("Colab has all six packages, so you should reach the same ", {}), ("UU-92-5650", AMs),
           (". Tick the step 6 boxes.", {})])
ap(right, [("Step 7 is the catch:", B), (" pushing from Colab is awkward. Use ", {}),
           ("File > Save a copy in GitHub", B), (", pointing at your own fork, and ", {}),
           ("tell us you are on Colab today", B),
           (" so we can make sure the labs work for you before they count.", {})])

# ---- footer
p = para(doc, before=7, after=0)
pr = p._p.get_or_add_pPr(); set_borders(pr, {"top": (RULE, "6", "4")}, "p")
run(p, "Before Thursday (Aug 27):", size=8.4, bold=True, color=MUTED)
run(p, " watch the data wrangling video, and finish this checklist if you got stuck — bring the "
       "questions you could not resolve. Thursday works in the same folder, in the same environment, "
       "on the same cars data, so an unfinished setup is the one thing that will put you behind.",
    size=8.4, color=MUTED)

_z = doc.settings.element.find(qn("w:zoom"))
if _z is not None and _z.get(qn("w:percent")) is None:
    _z.set(qn("w:percent"), "100")

doc.save("setup-checklist.docx")
print("saved setup-checklist.docx")
